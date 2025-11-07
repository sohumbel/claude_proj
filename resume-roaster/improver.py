"""Resume improvement module for generating improved versions."""

import re
from typing import Dict
import docx
from docx.shared import Pt, RGBColor
from config import BUZZWORDS, WEAK_VERBS, STRONG_VERBS


def generate_improved_resume(parsed_data: Dict, analysis: Dict, output_path: str) -> str:
    """
    Generate an improved resume document with fixes applied.

    Args:
        parsed_data: Parsed resume data
        analysis: Analysis results
        output_path: Path for the improved resume file

    Returns:
        Path to the generated improved resume
    """
    file_type = parsed_data.get("formatting", {}).get("file_type", "docx")

    if file_type == "pdf":
        # For PDF input, create DOCX output (can't easily edit PDFs)
        output_path = output_path.replace('.pdf', '_improved.docx')
        return create_improved_docx(parsed_data, analysis, output_path)
    else:
        output_path = output_path.replace('.docx', '_improved.docx')
        return create_improved_docx(parsed_data, analysis, output_path)


def create_improved_docx(parsed_data: Dict, analysis: Dict, output_path: str) -> str:
    """Create an improved DOCX resume."""
    doc = docx.Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add contact information
    contact = parsed_data.get("contact", {})
    if contact.get("name"):
        name_para = doc.add_paragraph(contact["name"])
        name_para.runs[0].font.size = Pt(16)
        name_para.runs[0].font.bold = True

    contact_info = []
    if contact.get("email"):
        contact_info.append(contact["email"])
    if contact.get("phone"):
        contact_info.append(contact["phone"])
    if contact.get("linkedin"):
        contact_info.append(contact["linkedin"])

    if contact_info:
        contact_para = doc.add_paragraph(" | ".join(contact_info))
        contact_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # Spacing

    # Add summary (improved if present)
    summary = parsed_data.get("summary")
    if summary:
        doc.add_heading("Professional Summary", level=2)
        improved_summary = improve_text(summary)
        doc.add_paragraph(improved_summary)
    else:
        # Add placeholder
        doc.add_heading("Professional Summary", level=2)
        placeholder = doc.add_paragraph("[ADD: 2-3 sentence professional summary highlighting your key qualifications and career goals]")
        placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        placeholder.runs[0].font.italic = True

    # Add experience section
    experience = parsed_data.get("experience", [])
    if experience:
        doc.add_heading("Work Experience", level=2)

        for exp in experience:
            # Company and title
            title_line = f"{exp.get('title', 'Position Title')} | {exp.get('company', 'Company Name')}"
            title_para = doc.add_paragraph(title_line)
            title_para.runs[0].font.bold = True

            # Dates
            if exp.get("dates"):
                date_para = doc.add_paragraph(exp["dates"])
                date_para.runs[0].font.italic = True
                date_para.runs[0].font.size = Pt(10)

            # Improved bullets
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                improved_bullet = improve_bullet(bullet, analysis)
                bullet_para = doc.add_paragraph(improved_bullet, style='List Bullet')

            doc.add_paragraph()  # Spacing between jobs

    # Add education section
    education = parsed_data.get("education", [])
    if education:
        doc.add_heading("Education", level=2)

        for edu in education:
            degree_line = edu.get("degree", "Degree")
            degree_para = doc.add_paragraph(degree_line)
            degree_para.runs[0].font.bold = True

            if edu.get("school"):
                school_para = doc.add_paragraph(edu["school"])

            date_info = []
            if edu.get("dates"):
                date_info.append(edu["dates"])
            if edu.get("gpa"):
                date_info.append(f"GPA: {edu['gpa']}")

            if date_info:
                info_para = doc.add_paragraph(" | ".join(date_info))
                info_para.runs[0].font.size = Pt(10)

            doc.add_paragraph()

    # Add skills section
    skills = parsed_data.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=2)
        skills_text = ", ".join(skills[:15])  # Limit to top 15 skills
        doc.add_paragraph(skills_text)

    # Add improvement notes at the end
    doc.add_page_break()
    doc.add_heading("📝 Improvement Notes", level=1)

    notes_para = doc.add_paragraph(
        "This is an improved version of your resume with the following changes applied:\n"
    )

    for i, fix in enumerate(analysis.get("critical_fixes", []), 1):
        doc.add_paragraph(f"{i}. {fix}", style='List Bullet')

    doc.add_paragraph("\nItems marked with [ADD:...] or [PLACEHOLDER] require your input with specific details.\n")

    # Save the document
    try:
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise ValueError(f"Failed to save improved resume: {str(e)}")


def improve_text(text: str) -> str:
    """Improve general text by removing buzzwords."""
    improved = text

    for buzzword in BUZZWORDS:
        pattern = re.compile(re.escape(buzzword), re.IGNORECASE)
        improved = pattern.sub("[REPLACE WITH SPECIFIC ACHIEVEMENT]", improved)

    return improved


def improve_bullet(bullet: str, analysis: Dict) -> str:
    """Improve a single bullet point."""
    improved = bullet

    # Replace weak verbs with strong ones
    for weak_verb in WEAK_VERBS:
        if weak_verb.lower() in improved.lower():
            # Get a random strong verb
            suggestion = STRONG_VERBS[0]  # Could randomize
            pattern = re.compile(re.escape(weak_verb), re.IGNORECASE)
            improved = pattern.sub(suggestion, improved, count=1)
            break

    # Check if bullet lacks metrics
    has_metrics = bool(re.search(r'\d+[%$]?|\$\d+', improved))
    if not has_metrics:
        # Add placeholder for metrics
        improved += " [ADD: Quantify the impact - how much, how many, percentage increase/decrease]"

    # Replace buzzwords
    for buzzword in BUZZWORDS:
        pattern = re.compile(re.escape(buzzword), re.IGNORECASE)
        if pattern.search(improved):
            improved = pattern.sub("[SPECIFIC ACHIEVEMENT]", improved)

    return improved


def suggest_improvements(parsed_data: Dict, analysis: Dict) -> Dict:
    """
    Generate improvement suggestions without creating a file.

    Args:
        parsed_data: Parsed resume data
        analysis: Analysis results

    Returns:
        Dictionary with improvement suggestions
    """
    suggestions = {
        "contact_improvements": [],
        "summary_improvements": [],
        "experience_improvements": [],
        "education_improvements": [],
        "skills_improvements": [],
        "formatting_improvements": []
    }

    # Contact improvements
    contact = parsed_data.get("contact", {})
    if not contact.get("email"):
        suggestions["contact_improvements"].append("Add professional email address")
    if not contact.get("linkedin"):
        suggestions["contact_improvements"].append("Add LinkedIn profile URL")

    # Summary improvements
    if not parsed_data.get("summary"):
        suggestions["summary_improvements"].append(
            "Add a professional summary (2-3 sentences highlighting your value proposition)"
        )

    # Experience improvements
    experience = parsed_data.get("experience", [])
    for i, exp in enumerate(experience):
        exp_suggestions = []

        for bullet in exp.get("bullets", []):
            # Check for weak verbs
            has_weak_verb = any(weak in bullet.lower() for weak in WEAK_VERBS)
            if has_weak_verb:
                exp_suggestions.append(f"Bullet '{bullet[:50]}...' uses weak language - use action verbs")

            # Check for metrics
            has_metrics = bool(re.search(r'\d+', bullet))
            if not has_metrics:
                exp_suggestions.append(f"Bullet '{bullet[:50]}...' lacks quantifiable results - add numbers")

        if exp_suggestions:
            suggestions["experience_improvements"].append({
                "job": f"{exp.get('title')} at {exp.get('company')}",
                "suggestions": exp_suggestions
            })

    # Formatting improvements
    formatting = parsed_data.get("formatting", {})
    if formatting.get("page_count", 0) > 2:
        suggestions["formatting_improvements"].append("Reduce resume to 1-2 pages")

    return suggestions
